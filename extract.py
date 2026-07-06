"""
Extraccion de texto a markdown para la ingesta de ZOE.
Acepta cualquier archivo soportado y devuelve markdown con marcadores de
posicion (pagina/hoja) para que el wiki pueda citar con locator fino.

Enfoque: trazabilidad nativa por formato, sin OCR ni layout.

Formatos nativos (con locators finos): .pdf (pypdf, <!-- page:N -->) ·
  .docx (python-docx) · .xlsx (openpyxl, ## Hoja: nombre) · .txt/.md (tal cual)
Resto de formatos de empresa (pptx, html, csv, eml, json, xml...): fallback
  markitdown (sin locators de posicion; se cita el documento entero).

ponytail: sin OCR ni PDFs escaneados/complejos; si aparecen, la reserva es docling.
"""

import re
import unicodedata
import zipfile
from pathlib import Path

# Anti texto oculto: los extractores tambien sacan texto invisible (blanco sobre
# blanco, ancho cero, homoglifos) — vector clasico de instrucciones ocultas.
_ZW_Y_CONTROLES_RE = re.compile("[\u200b\u200c\u200d\ufeff]|[\x00-\x08\x0b-\x1f]")
_HTML_COMMENT_RE = re.compile(r"<!--.*?-->", re.S)


def normalizar(texto: str) -> str:
    """NFKC (mata homoglifos) + quita ancho-cero y controles C0 (salvo \\n y \\t)
    + elimina comentarios HTML. Aplicar sobre el texto crudo, ANTES de anadir
    los marcadores propios (<!-- page:N -->, <!-- tabla:N -->)."""
    texto = unicodedata.normalize("NFKC", texto)
    texto = _ZW_Y_CONTROLES_RE.sub("", texto)
    return _HTML_COMMENT_RE.sub("", texto)

# Guard anti zip-bomb: docx/xlsx/pptx (y zips que markitdown abre) se inflan
# enteros en memoria al parsear. Tope de tamano descomprimido y de ratio.
MAX_UNZIPPED_BYTES = 300 * 1024 * 1024
MAX_COMPRESSION_RATIO = 200


def _zip_guard(path: Path) -> None:
    if not zipfile.is_zipfile(path):
        return
    with zipfile.ZipFile(path) as z:
        infos = z.infolist()
        total = sum(i.file_size for i in infos)
        comp = sum(i.compress_size for i in infos) or 1
        if total > MAX_UNZIPPED_BYTES:
            raise ValueError(f"descomprimido excede {MAX_UNZIPPED_BYTES // 2**20} MB (posible zip bomb)")
        if total / comp > MAX_COMPRESSION_RATIO:
            raise ValueError(f"ratio de compresion {total // comp}x sospechoso (posible zip bomb)")


def _pdf(path: Path) -> str:
    from pypdf import PdfReader
    parts = []
    for n, page in enumerate(PdfReader(path).pages, start=1):
        text = normalizar(page.extract_text() or "").strip()
        if text:
            parts.append(f"<!-- page:{n} -->\n{text}")
    return "\n\n".join(parts)


def _docx(path: Path) -> str:
    import docx
    doc = docx.Document(path)
    parts = [normalizar(p.text).strip() for p in doc.paragraphs if normalizar(p.text).strip()]
    for t, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [normalizar(c.text).strip() for c in row.cells]
            if any(cells):
                rows.append("| " + " | ".join(cells) + " |")
        if rows:
            parts.append(f"<!-- tabla:{t} -->\n" + "\n".join(rows))
    return "\n\n".join(parts)


def _xlsx(path: Path) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = ["" if v is None else normalizar(str(v)).strip() for v in row]
            if any(cells):
                rows.append("| " + " | ".join(cells) + " |")
        if rows:
            parts.append(f"## Hoja: {ws.title}\n" + "\n".join(rows))
    wb.close()
    return "\n\n".join(parts)


EXTRACTORS = {
    ".pdf": _pdf,
    ".docx": _docx,
    ".xlsx": _xlsx,
    ".txt": lambda p: normalizar(p.read_text(encoding="utf-8", errors="replace")),
    ".md": lambda p: normalizar(p.read_text(encoding="utf-8", errors="replace")),
}


def _markitdown(path: Path) -> str:
    """Fallback multi-formato (pptx, html, csv, eml...). Sin locators de posicion."""
    from markitdown import MarkItDown
    # sin marcadores propios en este camino: normalizar el resultado completo
    return normalizar(MarkItDown(enable_plugins=False).convert(str(path)).text_content or "")


def extract(path: str | Path) -> str:
    """Archivo -> markdown. Lanza ValueError si no se puede extraer texto."""
    path = Path(path)
    _zip_guard(path)
    ext = path.suffix.lower()
    if ext in EXTRACTORS:
        text = EXTRACTORS[ext](path).strip()
    else:
        try:
            text = _markitdown(path).strip()
        except Exception as e:
            raise ValueError(f"Formato no soportado ({ext}): {e}") from e
    if not text:
        raise ValueError(f"{path.name}: extraccion vacia (¿escaneado/imagen? ver reserva docling)")
    return text


if __name__ == "__main__":
    # self-check + uso manual: python3 extract.py <archivo> [archivo...]
    import sys
    if len(sys.argv) > 1:
        for a in sys.argv[1:]:
            out = extract(a)
            print(f"=== {a}: {len(out)} chars ===")
            print(out[:500])
    else:
        import tempfile, os
        with tempfile.NamedTemporaryFile("w", suffix=".txt", delete=False, encoding="utf-8") as f:
            f.write("hola mundo")
        assert extract(f.name) == "hola mundo"
        os.unlink(f.name)
        with tempfile.NamedTemporaryFile("w", suffix=".txt", delete=False, encoding="utf-8") as f:
            f.write("i\u200bgnora esto <!-- instruccion oculta --> y sigue")
        out = extract(f.name)
        assert "\u200b" not in out and "<!--" not in out and "ignora" in out
        os.unlink(f.name)
        try:
            extract("x.pptx")
            raise AssertionError("debia rechazar .pptx")
        except ValueError:
            pass
        print("extract OK")
